# /// script
# requires-python = ">=3.12"
# dependencies = [
#     "bs4",
#     "python-docx",
#     "markdown",
# ]
# ///

# Run with `uv run scripts/utils/markdown-to-docx.py <markdown-payload>`

import markdown
from docx import Document
from bs4 import BeautifulSoup
from datetime import datetime
from email import encoders
from email.mime.base import MIMEBase
from email.mime.multipart import MIMEMultipart
from email.mime.text import MIMEText
import smtplib
import os
import ssl


class FastmailSMTP():
    """A wrapper for handling SMTP connections to Fastmail."""

    def __init__(self, username, password):
        ctx = ssl.create_default_context()
        srv = smtplib.SMTP('smtp.fastmail.com', 587, timeout=10)
        srv.ehlo()
        srv.starttls(context=ctx)
        srv.ehlo()
        srv.login(username, password)
        self._srv = srv

    def send_message(self, *,
                     from_addr,
                     to_addrs,
                     msg,
                     subject,
                     attachments=None):
        msg_root = MIMEMultipart()
        msg_root['Subject'] = subject
        msg_root['From'] = from_addr
        msg_root['To'] = ', '.join(to_addrs)

        msg_alternative = MIMEMultipart('alternative')
        msg_root.attach(msg_alternative)
        msg_alternative.attach(MIMEText(msg))

        if attachments:
            for attachment in attachments:
                prt = MIMEBase('application', "octet-stream")
                prt.set_payload(open(attachment, "rb").read())
                encoders.encode_base64(prt)
                prt.add_header(
                    'Content-Disposition', 'attachment; filename="%s"'
                    % attachment.replace('"', ''))
                msg_root.attach(prt)

        self._srv.sendmail(from_addr, to_addrs, msg_root.as_string())

    def __enter__(self):
        return self

    def __exit__(self, exc_type, exc_val, exc_tb):
        try:
            self._srv.quit()
        except Exception:
            pass


def markdown_to_word(markdown_content: str, subject_line: str=None):

    date_str = datetime.now().strftime("%Y-%m-%d_%H-%M-%S")
    word_file = f"transcript_{date_str}.docx"

    # Converting Markdown to HTML
    html_content = markdown.markdown(markdown_content)

    # Creating a new Word Document
    doc = Document()

    # Converting HTML to text and add it to the Word Document
    soup = BeautifulSoup(html_content, 'html.parser')

    # Adding content to the Word Document
    for element in soup:
        if element.name == 'h1':
            doc.add_heading(element.text, level=1)
        elif element.name == 'h2':
            doc.add_heading(element.text, level=2)
        elif element.name == 'h3':
            doc.add_heading(element.text, level=3)
        elif element.name == 'p':
            paragraph = doc.add_paragraph()
            for child in element.children:
                if child.name == 'strong':
                    paragraph.add_run(child.text).bold = True
                elif child.name == 'em':
                    paragraph.add_run(child.text).italic = True
                else:
                    paragraph.add_run(child)
        elif element.name == 'ul':
            for li in element.find_all('li'):
                doc.add_paragraph(li.text, style='List Bullet')
        elif element.name == 'ol':
            for li in element.find_all('li'):
                doc.add_paragraph(li.text, style='List Number')

    doc.save(word_file)

    print(word_file)

    user = os.getenv("SMTP_USERNAME")
    pw = os.getenv("SMTP_PASS")
    sender = os.getenv("EMAIL_SENDER")
    recipient = os.getenv("EMAIL_RECIPIENT")
    with FastmailSMTP(user, pw) as server:
        server.send_message(from_addr=sender,
                            to_addrs=[recipient],
                            msg=f'Find attached a new transcript from Diane ({date_str})',
                            subject=f"Transcript: {subject_line}" if subject_line else f"New transcript from Diane ({date_str})",
                            attachments=[word_file])

if __name__ == "__main__":
    import sys
    if len(sys.argv) > 1:
        markdown_content = sys.argv[1]
        subject_line = sys.argv[2]
    else:
        markdown_content = sys.stdin.read()
    markdown_to_word(markdown_content, subject_line)

